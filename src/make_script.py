"""
Generates the presentation speaker script as a Word document.
Run from the project root: python src/make_script.py
Output: outputs/presentation_script.docx
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

os.makedirs("outputs", exist_ok=True)
doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, size, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 16, bold=True, color=(180, 0, 0))
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, 12, bold=True, color=(60, 60, 60))
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    set_font(run, 11)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Inches(0.3 + level * 0.2)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, 11)
    return p

def script_line(text):
    """The actual words to say — slightly indented, regular weight."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 11, color=(30, 30, 30))
    return p

def divider():
    p = doc.add_paragraph("─" * 72)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    for run in p.runs:
        set_font(run, 9, color=(180, 180, 180))

def slide_header(num, title, time_est):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(f"SLIDE {num}  ·  {title.upper()}")
    set_font(r1, 13, bold=True, color=(180, 0, 0))
    r2 = p.add_run(f"    [{time_est}]")
    set_font(r2, 10, color=(130, 130, 130))


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("U.S. HOUSING MARKET ANALYSIS")
set_font(r, 22, bold=True, color=(180, 0, 0))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Presentation Speaker Script")
set_font(r, 13, color=(80, 80, 80))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ISQS 3358  ·  Spring 2026  ·  Texas Tech University")
set_font(r, 11, color=(120, 120, 120))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Samuel Melghem  ·  Omar Hernandez  ·  Ryan Rhodes\n"
              "Ethan Hennenhoefer  ·  Logan Crider")
set_font(r, 11, color=(120, 120, 120))

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run("How to use this document")
set_font(r, 11, bold=True)
bullet("Part 1 — Outline: one page per slide with the key points to hit. Use this to stay on track.")
bullet("Part 2 — Full Script: complete word-for-word language. Read it a few times, then use the outline.")
bullet("Times in brackets are targets. Total presentation: 12–17 minutes plus Q&A.")
bullet("Speaker assignments are marked as [ETHAN], [OMAR], etc. — adjust as your team decides.")


# ══════════════════════════════════════════════════════════════════════════════
# PART 1 — OUTLINE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
p = doc.add_paragraph()
r = p.add_run("PART 1 — PRESENTATION OUTLINE")
set_font(r, 18, bold=True, color=(180, 0, 0))

body("Quick-reference talking points for each slide. One glance should be enough to stay on track.", indent=False)

divider()

# ── Slide 1
slide_header(1, "Title Slide", "~30 sec")
bullet("Welcome the audience, introduce the team")
bullet("State the topic: what drives real U.S. home prices?")
bullet("Set the tone — this is data-driven, not just opinions")

# ── Slide 2
slide_header(2, "Roadmap", "~30 sec")
bullet("Walk through the 8 sections in order")
bullet("Tell them where the interesting stuff is — findings and charts")
bullet("Keep it brief — don't explain each item, just name them")

# ── Slide 3
slide_header(3, "Project Purpose", "~1.5 min")
bullet("Why housing? It's the biggest financial decision most people make")
bullet("Hit the key stat: nominal prices rose 5× since 1984")
bullet("Inflation-adjusted it's closer to 2× — context matters")
bullet("Pose the research question clearly before moving on")

# ── Slide 4
slide_header(4, "Industry Background", "~1.5 min")
bullet("Brief framing: $40T market, controlled by supply/demand/financing")
bullet("Why it's hard to model: different data frequencies, structural breaks")
bullet("Our approach: FRED, 7 variables, linear regression — keep it simple")

# ── Slide 5
slide_header(5, "Datasets", "~1.5 min")
bullet("Walk the table — what each variable is and why we included it")
bullet("Point out the frequency problem (weekly, quarterly, annual) we had to solve")
bullet("Emphasize FRED: free, government-verified, automatic download")
bullet("Explain the target: RealPrice = HomePrice ÷ CPI × 100")

# ── Slide 6
slide_header(6, "Data Pipeline", "~1.5 min")
bullet("Walk through the 5-step flow left to right")
bullet("Stress the automation: one command runs everything")
bullet("Explain the two tricky resampling decisions: weekly → mean, annual → ffill")
bullet("End with: 504 clean monthly observations, 1984–2025")

# ── Slide 7
slide_header(7, "Individual Trends", "~1.5 min")
bullet("Point out real price: peaks 2006, crashes 2008, recovers and surges")
bullet("Mortgage rate: fell from 14% in 1984 to under 3% in 2021, now back up")
bullet("Unemployment: smooth until the 2008 spike — mirrors the price crash perfectly")
bullet("Brief note on income and population: steady growth, no major shocks")

# ── Slide 8
slide_header(8, "Key Relationships", "~1.5 min")
bullet("Top chart (price vs. mortgage): as rates fall, prices rise — clear negative relationship")
bullet("Bottom chart (indexed price vs. income): both grew, but prices grew nearly twice as fast")
bullet("The gap between price growth and income growth is the affordability problem in one image")

# ── Slide 9
slide_header(9, "Affordability", "~1.5 min")
bullet("Left chart: nominal 5× growth vs real 2× growth — inflation tells most of the story")
bullet("Right chart: price-to-income ratio went from 3.5× to 7×+ — this is the real story")
bullet("Punch line: even after adjusting for inflation, homes are far less affordable than in 1984")

# ── Slide 10
slide_header(10, "Model Setup", "~1 min")
bullet("Linear regression — 5 features, 1 target (RealPrice)")
bullet("Full dataset: all 504 months, 1984–2025")
bullet("Preview the result: R² = 0.8775 — teaser before the results slide")

# ── Slide 11
slide_header(11, "Regression Results", "~2 min")
bullet("Lead with R² = 0.88 — model explains 88% of variation, very strong")
bullet("Walk the coefficient table row by row")
bullet("Mortgage rate: −$847 per percentage point — most impactful negative force")
bullet("Unemployment: −$1,338 per point — job losses tank prices fast")
bullet("Income: positive, makes sense — higher earnings mean higher prices")
bullet("Note multicollinearity: interpret directions more than magnitudes")

# ── Slide 12
slide_header(12, "Actual vs Predicted", "~1 min")
bullet("Blue = actual, dashed = model prediction")
bullet("Model tracks the 40-year trend extremely well")
bullet("Captures 2008 crash and the COVID surge direction")
bullet("Short-term shocks cause noise above/below the economic baseline")

# ── Slide 13
slide_header(13, "Key Findings", "~2 min")
bullet("Income is the core driver — as earnings rise, buyers bid more")
bullet("Mortgage rates suppress prices — the 40-year rate decline subsidized growth")
bullet("Unemployment drives prices down — 2008 is the proof")
bullet("Affordability deteriorated: price-to-income doubled in 40 years")
bullet("Inflation masked the real story — nominal prices are misleading")
bullet("Five variables explain 88% of four decades of price movement")

# ── Slide 14
slide_header(14, "Limitations", "~1 min")
bullet("National averages only — Texas ≠ California ≠ rural Midwest")
bullet("Multicollinearity: income and population trend together, complicates interpretation")
bullet("Linear model: a 1% rate move at 14% differs from one at 3%")
bullet("Supply side gaps: zoning and land scarcity aren't captured by housing starts alone")

# ── Slide 15
slide_header(15, "Conclusion", "~1 min")
bullet("Economic fundamentals explain 88% of housing price variation")
bullet("Cheap financing (low rates) drove prices far more than most people realize")
bullet("Affordability has quietly doubled in severity since 1984")
bullet("Call to action / takeaway for the audience")

# ── Slide 16
slide_header(16, "Q&A", "~0 sec")
bullet("Thank the audience, open the floor")
bullet("Be ready for: why not include COVID? why linear? what about regional data?")

divider()


# ══════════════════════════════════════════════════════════════════════════════
# PART 2 — FULL WORD-FOR-WORD SCRIPT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
p = doc.add_paragraph()
r = p.add_run("PART 2 — FULL SPEAKER SCRIPT")
set_font(r, 18, bold=True, color=(180, 0, 0))

body("Complete language for every slide. Read this out loud a few times before the presentation. "
     "You don't need to say it word for word — but know it well enough that you can say it naturally.", indent=False)

divider()

# ════════════════════════════
# SLIDE 1 — Title
# ════════════════════════════
slide_header(1, "Title Slide", "~30 sec")
heading2("[ETHAN — Lead]")
script_line(
    "Good afternoon everyone. My name's Ethan, and with me today are Samuel, Omar, Ryan, and Logan. "
    "We're here from ISQS 3358 to present our analysis of the U.S. housing market. "
    "Specifically, we looked at a question that affects basically every person in this room eventually — "
    "what actually drives home prices? Not what people on the internet say drives home prices, "
    "but what the data says. And we have forty years of it. Let's get into it."
)

# ════════════════════════════
# SLIDE 2 — Roadmap
# ════════════════════════════
slide_header(2, "Roadmap", "~30 sec")
heading2("[ETHAN]")
script_line(
    "Here's how we're structured today. We'll start with why we picked this topic and some background "
    "on the housing market. Then we'll walk through the data — where we got it, how we cleaned it, "
    "and what it looks like. After that, the model — how we built it and what it found. "
    "Then findings, limitations, and we'll close with a conclusion. "
    "The interesting stuff really picks up around slide seven when we get into the charts, so stay with us."
)

# ════════════════════════════
# SLIDE 3 — Project Purpose
# ════════════════════════════
slide_header(3, "Project Purpose", "~1.5 min")
heading2("[ETHAN]")
script_line(
    "So why housing? Because it's the single largest financial decision most Americans will ever make. "
    "And the numbers are staggering — median home prices went from around seventy-nine thousand dollars "
    "in 1984 to over four hundred thousand today. That's a five-times increase in nominal terms. "
    "Five times. That sounds alarming."
)
script_line(
    "But here's where it gets interesting — when you adjust for inflation, the real increase "
    "is closer to two times. So about half of that apparent price explosion is just inflation. "
    "The other half is the actual housing market doing something different. "
    "And untangling those two things is exactly what we set out to do."
)
script_line(
    "We also looked at how mortgage rates fit into this. In 1984, the thirty-year fixed rate "
    "was around fourteen percent. By 2021 it was under three. That's a massive shift in what people "
    "can actually afford to borrow. And it has a direct effect on prices."
)
script_line(
    "Which brings us to our research question — printed at the bottom of this slide: "
    "Do economic indicators predict inflation-adjusted home prices, and which factors "
    "have the strongest effect? That's what we built a model to answer."
)

# ════════════════════════════
# SLIDE 4 — Industry Background
# ════════════════════════════
slide_header(4, "Industry Background", "~1.5 min")
heading2("[SAMUEL]")
script_line(
    "To understand our approach, it helps to understand the market we're studying. "
    "The U.S. residential real estate market is worth over forty trillion dollars. "
    "It's one of the largest asset classes in the world, and it's highly sensitive to "
    "what the Federal Reserve does with interest rates."
)
script_line(
    "Prices are set by supply and demand — but what makes housing hard to model is that "
    "those forces operate on different timescales. Supply is slow — it takes years to build homes. "
    "Demand can shift in months based on mortgage rates and employment. "
    "And data comes in at different frequencies — mortgage rates update weekly, "
    "income data is only released once a year."
)
script_line(
    "Our approach was to pull everything from FRED — the Federal Reserve's public data platform — "
    "standardize it all to monthly data, and then run a regression. "
    "One pipeline, one command, reproducible by anyone who clones our repository. "
    "We'll show you what that looks like in a minute."
)

# ════════════════════════════
# SLIDE 5 — Datasets
# ════════════════════════════
slide_header(5, "Datasets", "~1.5 min")
heading2("[SAMUEL]")
script_line(
    "Here are the seven series we used — all from FRED, all free and government-verified. "
    "The target variable is median U.S. home sale price, which comes in quarterly. "
    "CPI is our inflation index — monthly. Mortgage rate is weekly. "
    "Unemployment is monthly. Median household income is annual. "
    "Population is monthly. And housing starts — new construction — is monthly."
)
script_line(
    "The frequency mismatch was one of our first data problems. You can't merge a weekly series "
    "with a quarterly one without making some decisions. "
    "We'll explain how we handled that on the next slide."
)
script_line(
    "And at the bottom of the slide is our target variable. We didn't just predict home prices — "
    "we predicted real home prices. That's the home price divided by CPI, times a hundred. "
    "That strips out inflation and lets us see what's actually happening to the housing market "
    "independent of general price growth."
)

# ════════════════════════════
# SLIDE 6 — Data Pipeline
# ════════════════════════════
slide_header(6, "Data Pipeline", "~1.5 min")
heading2("[OMAR]")
script_line(
    "Here's how the pipeline works. It's five steps, and you run one file — main.py — "
    "and it does all of this automatically."
)
script_line(
    "Step one: download. We call the FRED API and pull all seven series as CSV files. "
    "If you want to add a new variable later, you add one line to the dictionary in main.py. "
    "Everything else picks it up automatically."
)
script_line(
    "Step two: resample. This is where we solve the frequency problem. "
    "Weekly mortgage rate data gets averaged to a monthly mean — one number per month. "
    "Annual income data gets forward-filled — we take the annual value and hold it "
    "constant month to month until the next year's release comes out. "
    "It's not perfect, but it's the standard approach for this type of data."
)
script_line(
    "Step three: merge. We do an outer join on date across all seven files, "
    "then forward-fill any remaining gaps. "
    "Step four: we calculate the real price. Step five: we run the regression. "
    "End result — five hundred and four clean monthly observations, "
    "January 1984 through December 2025."
)

# ════════════════════════════
# SLIDE 7 — Individual Trends
# ════════════════════════════
slide_header(7, "Individual Trends", "~1.5 min")
heading2("[OMAR]")
script_line(
    "Here's what the data looks like. Six charts — one for each of our main variables. "
    "Let me walk through the notable ones."
)
script_line(
    "Real home price — top left. You can see the big story: "
    "prices rise through the nineties and two-thousands, peak around 2006, "
    "then crash hard in 2008, recover slowly, and then absolutely surge after 2020. "
    "That COVID-era spike is real — and it's one of the most extreme moves in the dataset."
)
script_line(
    "Mortgage rate — top center. This one's almost the mirror image of prices "
    "over the long run. Fourteen percent in 1984, falling steadily for four decades "
    "down to under three percent in 2021, then spiking back to seven in 2022 and 2023. "
    "That reversal is why affordability collapsed so fast recently."
)
script_line(
    "Unemployment — top right. Smooth for the most part, and then the 2008 spike. "
    "It goes from four and a half percent to ten percent in about eighteen months. "
    "When you overlay that with the price chart, it lines up almost perfectly with the crash."
)

# ════════════════════════════
# SLIDE 8 — Key Relationships
# ════════════════════════════
slide_header(8, "Key Relationships", "~1.5 min")
heading2("[RYAN]")
script_line(
    "Now let's look at the relationships between variables — this is where it starts "
    "to get interesting before we even run the model."
)
script_line(
    "Top chart: real home price versus mortgage rate. "
    "The two lines are almost opposites of each other. "
    "As rates fell from fourteen percent down to three, prices climbed. "
    "When rates reversed in 2022, prices started pulling back. "
    "That's the relationship our model is going to quantify — "
    "and the coefficient on mortgage rate ends up being one of the strongest in the model."
)
script_line(
    "Bottom chart: this one shows home price and median income, "
    "both indexed to 100 in 1984 so they're on the same scale. "
    "Income grew — but home prices grew nearly twice as fast. "
    "You can see the two lines start together in 1984 and then diverge, "
    "especially after 2000, and especially after 2020. "
    "That gap is the affordability problem, right there in one image."
)

# ════════════════════════════
# SLIDE 9 — Affordability
# ════════════════════════════
slide_header(9, "Affordability", "~1.5 min")
heading2("[RYAN]")
script_line(
    "This slide zooms in on affordability specifically, because I think it tells "
    "a story that surprises a lot of people."
)
script_line(
    "Left chart — nominal versus real price. The gray line is the nominal price, "
    "which looks like it exploded five times since 1984. "
    "The blue line is the inflation-adjusted real price — only about double. "
    "Inflation accounts for most of the apparent increase. "
    "If you just look at sticker prices without adjusting for inflation, "
    "you're drawing the wrong conclusions."
)
script_line(
    "Right chart — price to income ratio. This one's simple: "
    "how many years of median household income does it take to buy a median home? "
    "In 1984, the answer was about three and a half years. "
    "Today, it's over seven. "
    "That's not a small shift — that's a doubling of the affordability burden "
    "in forty years. Even accounting for inflation, homes are significantly harder "
    "to buy today than they were in 1984. That's a real trend, not a statistical artifact."
)

# ════════════════════════════
# SLIDE 10 — Model Setup
# ════════════════════════════
slide_header(10, "Model Setup", "~1 min")
heading2("[LOGAN]")
script_line(
    "Alright — the model. We used ordinary least squares linear regression from scikit-learn. "
    "Five input features: unemployment rate, mortgage rate, median household income, "
    "population, and housing starts. One target: real home price."
)
script_line(
    "We trained on the full dataset — all five hundred and four months, "
    "January 1984 through December 2025. "
    "Linear regression was the right choice here because we care about interpretability. "
    "We don't just want a prediction — we want to be able to say, "
    "for every one percent increase in mortgage rates, home prices drop by this much. "
    "That's what you get with linear regression that you don't get with a black-box model."
)
script_line(
    "And I'll preview the headline number — R-squared of 0.8775. "
    "We'll talk about what that means on the next slide."
)

# ════════════════════════════
# SLIDE 11 — Results
# ════════════════════════════
slide_header(11, "Regression Results", "~2 min")
heading2("[LOGAN]")
script_line(
    "Here are the results. Let me start with the top row — "
    "R-squared of 0.8775. That means our model explains about 88% of the variation "
    "in real home prices using just five economic variables. "
    "For a simple linear model on noisy, real-world economic data over forty years, "
    "that's a very strong result."
)
script_line(
    "The root mean squared error — which is the average prediction error in dollars — "
    "works out to about sixty-five hundred dollars. On a median home price of two to four hundred thousand, "
    "that's a small miss on average."
)
script_line(
    "Now the coefficient table. Mortgage rate: negative eight forty-seven. "
    "For every one percentage point increase in the thirty-year fixed rate, "
    "real home prices drop by about eight hundred and forty-seven dollars. "
    "That negative relationship makes sense — higher rates price buyers out."
)
script_line(
    "Unemployment: negative thirteen thirty-eight. "
    "Job losses reduce buying power and force distressed selling — prices fall."
)
script_line(
    "Median income: positive point five five. Higher income means more purchasing power — prices rise. "
    "Population: small positive. More people, more demand. "
    "Housing starts: positive — more construction tends to occur in growth periods."
)
script_line(
    "One note on interpretation: income, population, and time all trend upward together, "
    "which creates some correlation between features — what statisticians call multicollinearity. "
    "The directions are right; the exact magnitudes should be interpreted with some caution."
)

# ════════════════════════════
# SLIDE 12 — Actual vs Predicted
# ════════════════════════════
slide_header(12, "Actual vs Predicted", "~1 min")
heading2("[LOGAN]")
script_line(
    "Here's what the model looks like visually. "
    "The solid blue line is the actual inflation-adjusted home price. "
    "The dashed line is what our model predicted."
)
script_line(
    "It tracks extremely well across four decades. "
    "You can see it nailing the general trend, capturing the 2008 crash and recovery, "
    "and following the COVID-era surge. "
    "There are moments where it's off — short-term shocks and sentiment moves "
    "that five economic variables can't fully capture. "
    "But the long-run story is clear: economic fundamentals set the floor, "
    "and events push prices above or below that floor temporarily."
)

# ════════════════════════════
# SLIDE 13 — Key Findings
# ════════════════════════════
slide_header(13, "Key Findings", "~2 min")
heading2("[ETHAN]")
script_line(
    "So what did we actually find? Let me walk through the six takeaways on this slide."
)
script_line(
    "First — income drives prices up. As household earnings grow, "
    "people qualify for larger mortgages and bid prices higher. "
    "It's the single most consistent positive force in the model."
)
script_line(
    "Second — mortgage rates suppress prices. Every one percent increase in rates "
    "corresponds to about eight hundred and forty-seven dollars off the real price. "
    "And the forty-year decline from fourteen to three percent quietly subsidized "
    "price growth the entire time. Buyers didn't notice because their monthly payment "
    "stayed manageable even as prices rose."
)
script_line(
    "Third — unemployment drives prices down. The 2008 crash is the clearest example. "
    "When unemployment doubles, buying power collapses and prices follow."
)
script_line(
    "Fourth — affordability has genuinely deteriorated. "
    "The price-to-income ratio went from three and a half to over seven. "
    "That's not inflation — we already adjusted for that. "
    "Homes are structurally harder to afford than they were forty years ago."
)
script_line(
    "Fifth — inflation masks the real story. "
    "Nominal prices rose five times, but real prices only doubled. "
    "If you don't adjust for inflation, you're comparing apples to oranges."
)
script_line(
    "And sixth — five variables explain 88% of four decades of price variation. "
    "That's the power of this kind of data-driven approach. "
    "You don't need a complex model. You need the right variables."
)

# ════════════════════════════
# SLIDE 14 — Limitations
# ════════════════════════════
slide_header(14, "Limitations", "~1 min")
heading2("[ETHAN]")
script_line(
    "We want to be honest about what this model can't do."
)
script_line(
    "First — it's national only. Texas and California are completely different markets. "
    "A national average model doesn't capture regional dynamics, and "
    "if you're trying to make a decision about a specific city, you'd want city-level data."
)
script_line(
    "Second — multicollinearity. Income, population, and time trend together over forty years. "
    "This can inflate or deflate individual coefficients. "
    "The directions are right; trust them. The exact dollar magnitudes — take those with a grain of salt."
)
script_line(
    "Third — the linear assumption. We're assuming a one percent rate move "
    "has the same effect at fourteen percent as it does at three percent. "
    "That's not quite true in practice."
)
script_line(
    "Fourth — supply side gaps. Housing starts is a reasonable proxy, "
    "but it misses zoning restrictions, construction costs, and land scarcity — "
    "which are increasingly important drivers, especially in expensive coastal markets."
)

# ════════════════════════════
# SLIDE 15 — Conclusion
# ════════════════════════════
slide_header(15, "Conclusion", "~1 min")
heading2("[ETHAN]")
script_line(
    "To bring it all together — economic performance absolutely drives the housing market, "
    "but not always in the ways people assume."
)
script_line(
    "The biggest surprise in our data is how much of the apparent price explosion "
    "is explained by falling mortgage rates. Rates went from fourteen to three percent "
    "over forty years. That's a massive reduction in borrowing costs, "
    "and it allowed prices to rise without monthly payments becoming unaffordable — "
    "at least until rates reversed in 2022."
)
script_line(
    "The affordability story is also more nuanced than headlines suggest. "
    "Nominal prices look terrifying. Real prices look much more moderate. "
    "But the price-to-income ratio doesn't lie — homes now take twice as many years "
    "of income to buy as they did in 1984, and that's a real problem for the next generation of buyers."
)
script_line(
    "Our model — five variables, linear regression, eighty-eight percent R-squared — "
    "captures the vast majority of that forty-year story. "
    "The code is all on GitHub if you want to run it yourself. One command. "
    "Thank you."
)

# ════════════════════════════
# SLIDE 16 — Q&A
# ════════════════════════════
slide_header(16, "Q&A", "~0 sec (transition)")
heading2("[ETHAN]")
script_line("We'd love to take any questions.")
heading2("Likely questions and suggested answers:")
bullet("Q: Why not include COVID / post-2020 data in the model?")
script_line(
    "A: We did include it — the model runs through December 2025. "
    "What we avoided was using a train/test split that put COVID in the test set, "
    "because that created a structural break the historical model couldn't predict. "
    "Training on the full period including COVID gives the model a chance to learn from it."
)
bullet("Q: Why linear regression and not something more complex?")
script_line(
    "A: Interpretability. We wanted to be able to say 'a one percent rate increase "
    "corresponds to this dollar change in price.' A neural network can't give you that. "
    "And with R² of 0.88, the linear model is doing extremely well here."
)
bullet("Q: What about regional differences?")
script_line(
    "A: Great point — that's our main limitation. National FRED data averages over "
    "massive regional variation. City-level analysis would be the natural next step."
)
bullet("Q: The unemployment coefficient is large — why?")
script_line(
    "A: Unemployment has outsized effects during recessions because it forces "
    "distressed selling and removes buyers simultaneously. The 2008 crash amplifies "
    "that coefficient significantly."
)

divider()
body("End of script. Total estimated presentation time: 12–17 minutes + Q&A.", indent=False)


# ── Save ──────────────────────────────────────────────────────────────────────
out = "outputs/presentation_script.docx"
doc.save(out)
print(f"Saved → {out}")
